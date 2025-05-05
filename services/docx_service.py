"""
Улучшенная реализация DocxService с точной заменой буквенных плейсхолдеров
для каждого типа шаблона
"""
import random
from typing import Dict, Any, List, Tuple
from docx import Document
from docx.shared import Pt

from services.base_document_service import BaseDocumentService
from models.models import (
    Group, Student, Teacher, Discipline, ExamQuestion,
    ScheduleItem, Publication
)


class DocxService(BaseDocumentService):
    """
    Сервис для обработки DOCX-документов с точной заменой плейсхолдеров
    """

    TEMPLATE_TYPE_MASTER_TITLE = "master_title"
    TEMPLATE_TYPE_BACHELOR_TITLE = "bachelor_title"
    TEMPLATE_TYPE_EXAM_TICKET = "exam_ticket"
    TEMPLATE_TYPE_ABSTRACT = "abstract"
    TEMPLATE_TYPE_LAB_WORK = "lab_work"
    TEMPLATE_TYPE_COURSE_WORK = "course_work"
    TEMPLATE_TYPE_COURSE_PROJECT = "course_project"
    TEMPLATE_TYPE_PRACTICE_THEMES = "practice_themes"
    TEMPLATE_TYPE_PUBLICATIONS = "publications"
    TEMPLATE_TYPE_LITERATURE = "literature"
    TEMPLATE_TYPE_EXAM_QUESTIONS = "exam_questions"
    TEMPLATE_TYPE_TEACHER_SCHEDULE = "teacher_schedule"
    TEMPLATE_TYPE_SECTION_PROGRAM = "section_program"
    TEMPLATE_TYPE_PRACTICE_REPORT = "practice_report"
    TEMPLATE_TYPE_TASK = "task"
    TEMPLATE_TYPE_CLASSROOM_SCHEDULE = "classroom_schedule"
    TEMPLATE_TYPE_GENERIC = "generic"

    async def load_document(self, template_path: str) -> Document:
        """
        Загружает DOCX-документ из файла шаблона
        """
        return Document(template_path)

    async def save_document(self, document: Document, output_path: str) -> None:
        """
        Сохраняет DOCX-документ в файл
        """
        document.save(output_path)

    async def determine_template_type(self, template_id: int, template_name: str) -> str:
        """
        Определяет тип шаблона по его ID и имени
        """
        template_name_lower = template_name.lower()

        print(template_name_lower)

        if "титул маг" in template_name_lower:
            return self.TEMPLATE_TYPE_MASTER_TITLE
        elif "титул бак" in template_name_lower:
            return self.TEMPLATE_TYPE_BACHELOR_TITLE
        elif "билет" in template_name_lower:
            return self.TEMPLATE_TYPE_EXAM_TICKET
        elif "реферат" in template_name_lower:
            return self.TEMPLATE_TYPE_ABSTRACT
        elif "лабораторная работа" in template_name_lower:
            return self.TEMPLATE_TYPE_LAB_WORK
        elif "курсовая работа" in template_name_lower:
            return self.TEMPLATE_TYPE_COURSE_WORK
        elif "курсовой проект" in template_name_lower:
            return self.TEMPLATE_TYPE_COURSE_PROJECT
        elif "практических работ" in template_name_lower:
            return self.TEMPLATE_TYPE_PRACTICE_THEMES
        elif "публикаций" in template_name_lower:
            return self.TEMPLATE_TYPE_PUBLICATIONS
        elif "литератур" in template_name_lower:
            return self.TEMPLATE_TYPE_LITERATURE
        elif "вопрос" in template_name_lower and ("экзамен" in template_name_lower or "зачет" in template_name_lower):
            return self.TEMPLATE_TYPE_EXAM_QUESTIONS
        elif "расписание преподавателей" in template_name_lower:
            return self.TEMPLATE_TYPE_TEACHER_SCHEDULE
        elif "секции" in template_name_lower:
            return self.TEMPLATE_TYPE_SECTION_PROGRAM
        elif "практик" in template_name_lower and "отчет" in template_name_lower:
            return self.TEMPLATE_TYPE_PRACTICE_REPORT
        elif "задание" in template_name_lower:
            return self.TEMPLATE_TYPE_TASK
        elif "загруженность" in template_name_lower or 'загруженность аудиторий' == template_name_lower:
            return self.TEMPLATE_TYPE_CLASSROOM_SCHEDULE
        else:

            return self.TEMPLATE_TYPE_GENERIC

    async def find_highlighted_text(self, document: Document) -> List[Tuple[Any, str]]:
        """
        Находит все фрагменты текста с желтым выделением
        """
        highlighted_items = []

        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, 'font') and hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                    highlighted_items.append((run, run.text.strip()))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if hasattr(run, 'font') and hasattr(run.font,
                                                                'highlight_color') and run.font.highlight_color:
                                highlighted_items.append((run, run.text.strip()))

        return highlighted_items

    async def replace_highlighted_text(self, document: Document, replacement_map: Dict[str, str]) -> None:
        """
        Заменяет выделенный текст на соответствующие значения и убирает желтое выделение
        """
        highlighted_items = await self.find_highlighted_text(document)

        for run, text in highlighted_items:
            if text in replacement_map:

                run.text = str(replacement_map[text])

                run.font.highlight_color = None

            elif len(text) == 1 and text in replacement_map:

                run.text = str(replacement_map[text])

            run.font.highlight_color = None

    async def apply_times_new_roman(self, document: Document) -> None:
        """
        Устанавливает шрифт Times New Roman для всех элементов документа
        """

        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"

                if run.font.size:
                    continue

                run.font.size = Pt(12)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"

                            if run.font.size:
                                continue

                            run.font.size = Pt(12)

    async def process_master_title(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон титульного листа магистерской работы

        Шаблон титул маг - "на тему" -  J - "Фамилия Имя Отчество" - N
        Шифр P Группа - T
        Руководитель работы И. О. Фамилия - О
        """

        discipline_id = params.get('discipline_id')
        student_id = params.get('student_id')
        teacher_id = params.get('teacher_id')

        replacements = {}

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['J'] = "Исследование и разработка методов машинного обучения"
        else:
            replacements['J'] = "Исследование и разработка методов машинного обучения"

        if student_id:
            student = await Student.get(id=student_id).prefetch_related('group')
            replacements['N'] = student.full_name
            if student.group:
                replacements['T'] = student.group.code
            else:
                replacements['T'] = "КМБО-05-21"

        if teacher_id:
            teacher = await Teacher.get(id=teacher_id)

            name_parts = teacher.full_name.split()
            if len(name_parts) >= 3:

                initials = f"{name_parts[1][0]}.{name_parts[2][0]}."
                replacements['O'] = f"{initials} {name_parts[0]}"
            else:
                replacements['O'] = teacher.full_name

        await self.replace_highlighted_text(document, replacements)

    async def process_bachelor_title(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон титульного листа бакалаврской работы

        Шаблон титул бал - "на тему" -  J - "Фамилия Имя Отчество" - N
        Шифр P Группа - T
        Руководитель работы - И. О. Фамилия - О
        Консультант (*при наличии*) - *ученая степень, ученое звание, должность - U*
        И. О. Фамилия - M
        """

        discipline_id = params.get('discipline_id')
        student_id = params.get('student_id')
        teacher_id = params.get('teacher_id')

        replacements = {}

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['J'] = "Разработка программного обеспечения"
        else:
            replacements['J'] = "Разработка программного обеспечения"

        if student_id:
            student = await Student.get(id=student_id).prefetch_related('group')
            replacements['N'] = student.full_name
            if student.group:
                replacements['T'] = student.group.code
            else:
                replacements['T'] = "КМБО-05-21"

        if teacher_id:
            teacher = await Teacher.get(id=teacher_id)

            name_parts = teacher.full_name.split()
            if len(name_parts) >= 3:
                initials = f"{name_parts[1][0]}.{name_parts[2][0]}."
                replacements['O'] = f"{initials} {name_parts[0]}"
                replacements['M'] = f"{initials} {name_parts[0]}"
            else:
                replacements['O'] = teacher.full_name
                replacements['M'] = teacher.full_name

            replacements['U'] = "к.т.н., доцент"
        else:

            replacements['O'] = "И.О. Фамилия"
            replacements['M'] = "И.О. Фамилия"
            replacements['U'] = "к.т.н., доцент"

        await self.replace_highlighted_text(document, replacements)

    async def process_exam_ticket(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон экзаменационного билета
        """
        discipline_id = params.get('discipline_id')
        ticket_number = params.get('ticket_number', random.randint(1, 30))

        replacements = {
            'K': str(ticket_number),
            'T': "Программирование и алгоритмы",
            'N': "1"
        }

        example_questions = [
            "Основные принципы программирования",
            "Алгоритмы сортировки и их сложность",
            "Методы оптимизации программного кода"
        ]

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['T'] = discipline.name

            questions = await ExamQuestion.filter(discipline_id=discipline_id).order_by('number')

            if questions:
                selected_questions = questions[:3] if len(questions) <= 3 else random.sample(list(questions), 3)
                example_questions = [q.text for q in selected_questions]

        success = await self.replace_exam_ticket_questions(document, example_questions)

        if not success:
            print("Не удалось найти плейсхолдеры вопросов в таблице, пробуем другой метод...")

            await self.replace_exam_ticket_questions(document, example_questions)

        await self.replace_highlighted_text(document, replacements)

        for paragraph in document.paragraphs:
            if "БИЛЕТ" in paragraph.text or "Билет" in paragraph.text:
                paragraph.text = paragraph.text.replace("№", f"№{ticket_number}")
                paragraph.text = paragraph.text.replace("БИЛЕТ", f"БИЛЕТ №{ticket_number}")

    async def process_abstract(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон титульного листа реферата

        по дисциплине N
        на тему: M
        Обучающийся: G
        Руководитель: P
        """
        discipline_id = params.get('discipline_id')
        student_id = params.get('student_id')
        teacher_id = params.get('teacher_id')

        replacements = {}

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['N'] = discipline.name
        else:
            replacements['N'] = "Программирование и алгоритмы"

        if student_id:
            student = await Student.get(id=student_id)
            replacements['G'] = student.full_name
        else:
            replacements['G'] = "Иванов Иван Иванович"

        if teacher_id:
            teacher = await Teacher.get(id=teacher_id)
            replacements['P'] = teacher.full_name
        else:
            replacements['P'] = "Петров Петр Петрович"

        replacements['M'] = "Современные методы анализа данных"

        await self.replace_highlighted_text(document, replacements)

    async def find_table_cells_with_content(self, document: Document, text_patterns: List[str]) -> List[
        Tuple[Any, str]]:
        """
        Находит ячейки таблицы, содержащие указанные паттерны текста

        Args:
            document: Документ Word
            text_patterns: Список паттернов текста для поиска

        Returns:
            Список пар (ячейка, найденный_паттерн)
        """
        result = []

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()

                    for pattern in text_patterns:
                        if pattern in cell_text:
                            result.append((cell, pattern))
                            break

        return result

    async def replace_cell_content(self, cell, pattern: str, replacement: str) -> None:
        """
        Заменяет содержимое ячейки таблицы

        Args:
            cell: Ячейка таблицы
            pattern: Текст для замены
            replacement: Новый текст
        """

        for paragraph in cell.paragraphs:

            if pattern in paragraph.text:
                alignment = paragraph.alignment

                new_text = paragraph.text.replace(pattern, replacement)

                paragraph.clear()
                run = paragraph.add_run(new_text)

                run.font.name = "Times New Roman"

                paragraph.alignment = alignment

    async def process_lab_work(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон титульного листа лабораторной работы

        по дисциплине N
        на тему: M
        Обучающийся: G
        Группа: S
        Руководитель: P
        """
        discipline_id = params.get('discipline_id')
        student_id = params.get('student_id')
        teacher_id = params.get('teacher_id')

        replacements = {}

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['N'] = discipline.name
        else:
            replacements['N'] = "Программирование и алгоритмы"

        if student_id:
            student = await Student.get(id=student_id).prefetch_related('group')
            replacements['G'] = student.full_name
            if student.group:
                replacements['S'] = student.group.code
            else:
                replacements['S'] = "КМБО-05-21"
        else:
            replacements['G'] = "Иванов Иван Иванович"
            replacements['S'] = "КМБО-05-21"

        if teacher_id:
            teacher = await Teacher.get(id=teacher_id)
            replacements['P'] = teacher.full_name
        else:
            replacements['P'] = "Петров Петр Петрович"

        replacements['M'] = "Разработка алгоритма сортировки данных"

        await self.replace_highlighted_text(document, replacements)

    async def process_course_work(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон титульного листа курсовой работы
        (Аналогичен лабораторной работе)
        """
        await self.process_lab_work(document, params)

    async def process_course_project(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон титульного листа курсового проекта
        (Аналогичен лабораторной работе)
        """
        await self.process_lab_work(document, params)

    async def process_practice_themes(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон тем практических работ

        Предмет: G
        Темы практических работ
        Семестр: L
        N, N, N - сами темы
        """
        discipline_id = params.get('discipline_id')

        replacements = {
            'L': "1"
        }

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['G'] = discipline.name

            practice_themes = [
                "Разработка алгоритмов обработки данных",
                "Использование структур данных",
                "Оптимизация вычислительных процессов"
            ]

            highlighted_items = await self.find_highlighted_text(document)
            theme_index = 0

            for run, text in highlighted_items:
                if text == 'N' and theme_index < len(practice_themes):
                    run.text = practice_themes[theme_index]
                    theme_index += 1
        else:
            replacements['G'] = "Программирование и алгоритмы"

            practice_themes = [
                "Разработка алгоритмов обработки данных",
                "Использование структур данных",
                "Оптимизация вычислительных процессов"
            ]

            highlighted_items = await self.find_highlighted_text(document)
            theme_index = 0

            for run, text in highlighted_items:
                if text == 'N' and theme_index < len(practice_themes):
                    run.text = practice_themes[theme_index]
                    theme_index += 1

        await self.replace_highlighted_text(document, replacements)

    async def replace_exam_ticket_questions(self, document, questions):
        """
        Метод, предназначенный специально для замены вопросов в таблице билета

        Args:
            document: Документ Word
            questions: Список вопросов (минимум 3)
        """
        print("Начинаем поиск и замену вопросов в билете...")

        if len(questions) < 3:
            questions = questions + ["Вопрос по умолчанию"] * (3 - len(questions))

        found_placeholders = []

        table_index = 0
        for table in document.tables:
            print(f"Таблица #{table_index}:")
            row_index = 0
            for row in table.rows:
                print(f"  Строка #{row_index}:")
                cell_index = 0
                for cell in row.cells:
                    print(f"    Ячейка #{cell_index}: '{cell.text}'")
                    cell_index += 1
                row_index += 1
            table_index += 1

        for table in document.tables:
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    if cell_text.startswith("1.") or cell_text.startswith("2.") or cell_text.startswith("3."):
                        print(f"Найдена ячейка с номером: '{cell_text}'")

                        if "M" in cell_text:
                            print(f"Внутри ячейки есть M: '{cell_text}'")
                            found_placeholders.append((cell, row_index, cell_index))


                        elif cell_index + 1 < len(row.cells):
                            next_cell = row.cells[cell_index + 1]
                            next_text = next_cell.text.strip()
                            if next_text == "M":
                                print(f"M найдена в следующей ячейке после номера")
                                found_placeholders.append((next_cell, row_index, cell_index + 1))

        if not found_placeholders:
            print("Не найдены ячейки с номерами, ищем просто M...")
            for table in document.tables:
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        if cell.text.strip() == "M":
                            print(f"Найдена ячейка только с M")
                            found_placeholders.append((cell, row_index, cell_index))

        if not found_placeholders:
            for table in document.tables:
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            if "M" in paragraph.text:
                                found_placeholders.append((cell, row_index, cell_index, paragraph))

        question_runs = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text.strip()
                            if run_text == "M":
                                question_runs.append(run)

        if question_runs:
            for i, run in enumerate(question_runs[:3]):
                if i < len(questions):
                    run.text = questions[i]
                    if hasattr(run, 'font') and hasattr(run.font, 'highlight_color'):
                        run.font.highlight_color = None


        elif found_placeholders:
            for i, placeholder_info in enumerate(found_placeholders[:3]):
                if i < len(questions):
                    if len(placeholder_info) == 4:
                        cell, _, _, paragraph = placeholder_info
                        paragraph.text = paragraph.text.replace("M", questions[i])
                    else:
                        cell, _, _ = placeholder_info
                        cell.text = questions[i]

        return len(question_runs) > 0 or len(found_placeholders) > 0

    async def replace_date_time_placeholders(self, paragraph, replacements):
        """
        Заменяет плейсхолдеры даты и времени, сохраняя исходное форматирование

        Args:
            paragraph: Параграф с плейсхолдерами
            replacements: Словарь замен
        """

        full_text = paragraph.text

        patterns = {
            "XX-XX-XXXX, XX:XX": replacements.get('XX-XX-XXXX, XX:XX', '01-05-2025, 14:30'),
            "XX-XX-XXXX,XX:XX": replacements.get('XX-XX-XXXX, XX:XX', '01-05-2025, 14:30'),
            "XX:XX-XX:XX": replacements.get('XX:XX-XX:XX', '09:00-10:30'),
            "XX:XX – XX:XX": replacements.get('XX:XX – XX:XX', '09:00 – 10:30'),
            "XX-XX-XXXX": replacements.get('XX-XX-XXXX', '01-05-2025')
        }

        found_pattern = None
        for pattern in patterns:
            if pattern in full_text:
                found_pattern = pattern
                break

        if not found_pattern:
            return False

        original_runs = list(paragraph.runs)

        temp_text = ""
        pattern_start_idx = -1

        for i, run in enumerate(original_runs):
            temp_text += run.text
            if found_pattern in temp_text and pattern_start_idx == -1:
                pattern_start_idx = temp_text.find(found_pattern)

        if pattern_start_idx == -1:
            return False

        for run in paragraph.runs:
            run.text = ""

        new_text = full_text.replace(found_pattern, patterns[found_pattern])

        if original_runs:
            new_run = paragraph.add_run(new_text)

            first_run = original_runs[0]
            if hasattr(first_run, 'font'):

                if hasattr(first_run.font, 'name'):
                    new_run.font.name = first_run.font.name

                if hasattr(first_run.font, 'size') and first_run.font.size:
                    new_run.font.size = first_run.font.size

                if hasattr(first_run.font, 'bold'):
                    new_run.font.bold = first_run.font.bold

                if hasattr(first_run.font, 'italic'):
                    new_run.font.italic = first_run.font.italic

                if hasattr(first_run.font, 'underline'):
                    new_run.font.underline = first_run.font.underline

        return True

    async def process_publications(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон списка публикаций

        Группа G
        Студент T
        Тема J
        """
        student_id = params.get('student_id')

        replacements = {}

        if student_id:
            student = await Student.get(id=student_id).prefetch_related('group')
            replacements['T'] = student.full_name

            if student.group:
                replacements['G'] = student.group.code
            else:
                replacements['G'] = "КМБО-05-21"

            publications = await Publication.filter(student_id=student_id)

            if publications:

                replacements['J'] = publications[0].title
            else:
                replacements['J'] = "Исследование методов машинного обучения"
        else:

            replacements['T'] = "Иванов Иван Иванович"
            replacements['G'] = "КМБО-05-21"
            replacements['J'] = "Исследование методов машинного обучения"

        await self.replace_highlighted_text(document, replacements)

    async def process_literature(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон списка литературы

        Предмет G
        N, N - сами названия
        """
        discipline_id = params.get('discipline_id')

        replacements = {}

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['G'] = discipline.name

            literature_examples = [
                "Кормен Т., Лейзерсон Ч., Ривест Р., Штайн К. Алгоритмы: построение и анализ, 3-е изд. — М.: «Вильямс», 2013.",
                "Кнут Д. Э. Искусство программирования. Том 1. Основные алгоритмы, 3-е изд. — М.: «Вильямс», 2006."
            ]

            highlighted_items = await self.find_highlighted_text(document)
            lit_index = 0

            for run, text in highlighted_items:
                if text == 'N' and lit_index < len(literature_examples):
                    run.text = literature_examples[lit_index]
                    lit_index += 1
        else:
            replacements['G'] = "Программирование и алгоритмы"

            literature_examples = [
                "Кормен Т., Лейзерсон Ч., Ривест Р., Штайн К. Алгоритмы: построение и анализ, 3-е изд. — М.: «Вильямс», 2013.",
                "Кнут Д. Э. Искусство программирования. Том 1. Основные алгоритмы, 3-е изд. — М.: «Вильямс», 2006."
            ]

            highlighted_items = await self.find_highlighted_text(document)
            lit_index = 0

            for run, text in highlighted_items:
                if text == 'N' and lit_index < len(literature_examples):
                    run.text = literature_examples[lit_index]
                    lit_index += 1

        await self.replace_highlighted_text(document, replacements)

    async def replace_split_time_placeholder(self, paragraph, new_time_value):
        """
        Заменяет разбитый плейсхолдер времени на новое значение

        Args:
            paragraph: Параграф, содержащий разбитый плейсхолдер
            new_time_value: Новое значение времени (например, '09:00-10:30')
        """

        runs_to_modify = []
        for run in paragraph.runs:
            if 'X' in run.text or ':' in run.text or '-' in run.text:
                runs_to_modify.append(run)

        if len(runs_to_modify) < 3:
            return

        if len(runs_to_modify) == len(new_time_value):
            for i, run in enumerate(runs_to_modify):
                run.text = new_time_value[i]
                if hasattr(run, 'font') and hasattr(run.font, 'highlight_color'):
                    run.font.highlight_color = None


        else:

            if runs_to_modify:
                runs_to_modify[0].text = new_time_value
                if hasattr(runs_to_modify[0], 'font') and hasattr(runs_to_modify[0].font, 'highlight_color'):
                    runs_to_modify[0].font.highlight_color = None

                for run in runs_to_modify[1:]:
                    run.text = ""
                    if hasattr(run, 'font') and hasattr(run.font, 'highlight_color'):
                        run.font.highlight_color = None

    async def process_exam_questions(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон списка вопросов для экзамена или зачета

        Предмет G
        Семестр T
        Вопросы к зачету/экзамену
        N, N, N - сами вопросы
        """
        discipline_id = params.get('discipline_id')

        replacements = {
            'T': "1"
        }

        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            replacements['G'] = discipline.name

            questions = await ExamQuestion.filter(discipline_id=discipline_id).order_by('number')

            if questions:

                highlighted_items = await self.find_highlighted_text(document)
                q_index = 0

                for run, text in highlighted_items:
                    if text == 'N' and q_index < len(questions):
                        run.text = questions[q_index].text
                        q_index += 1
            else:

                example_questions = [
                    "Основные принципы программирования",
                    "Алгоритмы сортировки и их сложность",
                    "Методы оптимизации программного кода"
                ]

                highlighted_items = await self.find_highlighted_text(document)
                q_index = 0

                for run, text in highlighted_items:
                    if text == 'N' and q_index < len(example_questions):
                        run.text = example_questions[q_index]
                        q_index += 1
        else:
            replacements['G'] = "Программирование и алгоритмы"

            example_questions = [
                "Основные принципы программирования",
                "Алгоритмы сортировки и их сложность",
                "Методы оптимизации программного кода"
            ]

            highlighted_items = await self.find_highlighted_text(document)
            q_index = 0

            for run, text in highlighted_items:
                if text == 'N' and q_index < len(example_questions):
                    run.text = example_questions[q_index]
                    q_index += 1

        await self.replace_highlighted_text(document, replacements)

    async def process_teacher_schedule(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон расписания преподавателей

        ФИО - М
        ДАТА - XX-XX-XXXX
        ВРЕМЯ - XX:XX-XX:XX
        ДИСЦИПЛИНА - P
        """
        teacher_id = params.get('teacher_id')
        day_of_week = params.get('day_of_week')

        replacements = {
            'XX-XX-XXXX': '01-05-2025',
            'XX:XX-XX:XX': '09:00-10:30'
        }

        if teacher_id:
            teacher = await Teacher.get(id=teacher_id)
            replacements['М'] = teacher.full_name
            replacements['M'] = teacher.full_name

            query = ScheduleItem.filter(teacher_id=teacher_id).prefetch_related('discipline')
            if day_of_week:
                query = query.filter(day_of_week=day_of_week)

            schedule_items = await query.order_by('time_slot__number')

            if schedule_items:

                replacements['P'] = schedule_items[0].discipline.name

                if hasattr(schedule_items[0], 'time_slot') and schedule_items[0].time_slot:
                    time_slot = schedule_items[0].time_slot
                    replacements['XX:XX-XX:XX'] = f"{time_slot.start_time}-{time_slot.end_time}"
            else:
                replacements['P'] = "Программирование и алгоритмы"
        else:
            replacements['М'] = "Иванов Иван Иванович"
            replacements['M'] = "Иванов Иван Иванович"
            replacements['P'] = "Программирование и алгоритмы"

        for paragraph in document.paragraphs:

            has_time_components = False
            x_count = 0

            for run in paragraph.runs:
                if 'X' in run.text:
                    x_count += run.text.count('X')
                    has_time_components = True

            if has_time_components and x_count >= 6:
                await self.replace_split_time_placeholder(paragraph, replacements['XX:XX-XX:XX'])

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:

                        has_time_components = False
                        x_count = 0

                        for run in paragraph.runs:
                            if 'X' in run.text:
                                x_count += run.text.count('X')
                                has_time_components = True

                        if has_time_components and x_count >= 6:
                            await self.replace_split_time_placeholder(paragraph, replacements['XX:XX-XX:XX'])

        await self.replace_highlighted_text(document, replacements)

    async def process_section_program(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон программы секции

        Дата и время: XX-XX-XXXX, XX:XX
        """
        replacements = {
            'XX-XX-XXXX, XX:XX': '10-05-2025, 14:30'
        }

        for paragraph in document.paragraphs:
            await self.replace_date_time_placeholders(paragraph, replacements)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        await self.replace_date_time_placeholders(paragraph, replacements)

        await self.replace_highlighted_text(document, replacements)

    async def process_practice_report(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон отчета по практике

        U СЕМЕСТР
        Группа - S
        Студент - M
        Руководитель практики - B
        """
        student_id = params.get('student_id')
        teacher_id = params.get('teacher_id')

        replacements = {
            'U': '7'
        }

        if student_id:
            student = await Student.get(id=student_id).prefetch_related('group')
            replacements['M'] = student.full_name

            if student.group:
                replacements['S'] = student.group.code
            else:
                replacements['S'] = "КМБО-05-21"
        else:
            replacements['M'] = "Иванов Иван Иванович"
            replacements['S'] = "КМБО-05-21"

        if teacher_id:
            teacher = await Teacher.get(id=teacher_id)
            replacements['B'] = teacher.full_name
        else:
            replacements['B'] = "Петров Петр Петрович"

        await self.replace_highlighted_text(document, replacements)

    async def process_task(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон задания для курса

        Обучающийся: N
        Шифр: O
        Группа: M
        Тема выпускной квалификационной работы - G
        """
        student_id = params.get('student_id')

        replacements = {}

        if student_id:
            student = await Student.get(id=student_id).prefetch_related('group')
            replacements['N'] = student.full_name
            replacements['O'] = "Александрова Виктория Петровна"

            if student.group:
                replacements['M'] = student.group.code
            else:
                replacements['M'] = "КМБО-05-21"

            replacements['G'] = "Разработка информационной системы управления образовательным процессом"
        else:

            replacements['N'] = "Иванов Иван Иванович"
            replacements['O'] = "Александрова Виктория Петровна"
            replacements['M'] = "КМБО-05-21"
            replacements['G'] = "Разработка информационной системы управления образовательным процессом"

        await self.replace_highlighted_text(document, replacements)

    async def process_classroom_schedule(self, document: Document, params: Dict[str, Any]) -> None:
        """
        Обрабатывает шаблон загруженности аудиторий с поддержкой разбитых плейсхолдеров времени
        """
        classroom_id = params.get('classroom_id')
        group_id = params.get('group_id')
        day_of_week = params.get('day_of_week')

        group_id_2 = params.get('group_id_2', group_id)

        group_name_1 = "КМБО-05-21"
        group_name_2 = "КМБО-02-21"

        if group_id:
            try:
                group = await Group.get(id=group_id)
                group_name_1 = group.code
            except:
                pass

        if group_id_2:
            try:
                group_2 = await Group.get(id=group_id_2)
                group_name_2 = group_2.code
            except:
                pass

        time_format_1 = "09:00 – 10:30"
        time_format_2 = "10:40 – 12:10"
        date_format = "01-05-2025"

        if classroom_id:
            query = ScheduleItem.filter(classroom_id=classroom_id).prefetch_related(
                'time_slot', 'discipline', 'group', 'teacher'
            )

            if day_of_week:
                query = query.filter(day_of_week=day_of_week)

            schedule_items = await query.order_by('time_slot__number')

            if schedule_items and len(schedule_items) > 0:

                group_name_1 = schedule_items[0].group.code

                if hasattr(schedule_items[0], 'time_slot') and schedule_items[0].time_slot:
                    time_slot_1 = schedule_items[0].time_slot
                    time_format_1 = f"{time_slot_1.start_time} – {time_slot_1.end_time}"

                if len(schedule_items) > 1:
                    group_name_2 = schedule_items[1].group.code
                    if hasattr(schedule_items[1], 'time_slot') and schedule_items[1].time_slot:
                        time_slot_2 = schedule_items[1].time_slot
                        time_format_2 = f"{time_slot_2.start_time} – {time_slot_2.end_time}"

        time_cells = []
        cell_index = 0

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:

                    cell_text = cell.text
                    if "XX:XX" in cell_text or "XX-XX" in cell_text or "XX:" in cell_text:
                        time_cells.append(cell)

        for i, cell in enumerate(time_cells):
            time_value = time_format_1 if i % 2 == 0 else time_format_2

            for paragraph in cell.paragraphs:

                has_x = False
                runs_with_x = []

                for run in paragraph.runs:
                    if 'X' in run.text:
                        has_x = True
                        runs_with_x.append(run)

                if has_x:

                    if len(runs_with_x) == 1 and ("XX:XX" in runs_with_x[0].text):
                        runs_with_x[0].text = runs_with_x[0].text.replace("XX:XX", time_value.split(' – ')[0])
                        if hasattr(runs_with_x[0], 'font') and hasattr(runs_with_x[0].font, 'highlight_color'):
                            runs_with_x[0].font.highlight_color = None


                    elif "XX:XX" in paragraph.text:

                        paragraph.text = paragraph.text.replace("XX:XX – XX:XX", time_value)
                        paragraph.text = paragraph.text.replace("XX:XX-XX:XX", time_value.replace(" – ", "-"))


                    else:

                        for run in runs_with_x:
                            if run.text == "XX:XX":
                                run.text = time_value.split(' – ')[0]
                            elif run.text == "XX:XX – XX:XX":
                                run.text = time_value
                            elif run.text == "XX:XX-XX:XX":
                                run.text = time_value.replace(" – ", "-")
                            else:

                                new_text = ""
                                for char in run.text:
                                    if char == 'X':

                                        if time_value:
                                            new_text += time_value[0]
                                            time_value = time_value[1:]
                                        else:
                                            new_text += '0'
                                    else:
                                        new_text += char
                                run.text = new_text

                            if hasattr(run, 'font') and hasattr(run.font, 'highlight_color'):
                                run.font.highlight_color = None

        highlighted_group_cells = []

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:

                            if hasattr(run, 'font') and hasattr(run.font,
                                                                'highlight_color') and run.font.highlight_color:
                                if run.text.strip() == 'M' or run.text.strip() == 'М':
                                    highlighted_group_cells.append((cell, run))

        for i, (cell, run) in enumerate(highlighted_group_cells):
            if i % 2 == 0:
                run.text = group_name_1
            else:
                run.text = group_name_2

            run.font.highlight_color = None

        for paragraph in document.paragraphs:
            if "Загруженность аудиторий на " in paragraph.text:

                date_placeholder = "XX-XX-XXXX"
                if date_placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(date_placeholder, date_format)

    async def generic_default_replacing(self, document, params):
        self.replacements = {}

        group_id = params.get('group_id')
        if group_id:
            group = await Group.get(id=group_id).prefetch_related('students')
            self.replacements.update({
                "group": group.code,
                "course": group.course,
                "students": "\n".join([f"{student.full_name} ({student.email})" for student in group.students]),
                "S": group.code,
                "T": group.code,
                "G": group.code
            })

        student_id = params.get('student_id')
        if student_id:
            student = await Student.get(id=student_id).prefetch_related('group')
            self.replacements.update({
                "student_name": student.full_name,
                "student_email": student.email,
                "G": student.full_name,
                "M": student.full_name,
                "N": student.full_name
            })

            if student.group:
                self.replacements.update({
                    "group": student.group.code,
                    "S": student.group.code,
                    "T": student.group.code
                })

        teacher_id = params.get('teacher_id')
        if teacher_id:
            teacher = await Teacher.get(id=teacher_id)
            self.replacements.update({
                "teacher_name": teacher.full_name,
                "teacher_email": teacher.email,
                "P": teacher.full_name,
                "B": teacher.full_name,
                "O": teacher.full_name,
                "U": "к.т.н., доцент"
            })

        discipline_id = params.get('discipline_id')
        if discipline_id:
            discipline = await Discipline.get(id=discipline_id)
            self.replacements.update({
                "discipline": discipline.name,
                "N": discipline.name
            })

        self.replacements.update({
            "J": "Исследование методов и алгоритмов",
            "L": "1",
            "К": "1",
            "XX-XX-XXXX": "01-05-2025",
            "XX:XX-XX:XX": "09:00-10:30"
        })

        await self.replace_highlighted_text(document, self.replacements)

        for paragraph in document.paragraphs:
            if any(placeholder in paragraph.text for placeholder in self.replacements.keys()) or any(
                    "{{" + placeholder + "}}" in paragraph.text for placeholder in self.replacements.keys()):
                alignment = paragraph.alignment

                new_text = await self.replace_placeholders_in_text(paragraph.text, self.replacements)

                paragraph.clear()
                run = paragraph.add_run(new_text)

                run.font.name = "Times New Roman"

                paragraph.alignment = alignment

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if any(placeholder in paragraph.text for placeholder in self.replacements.keys()) or any(
                                "{{" + placeholder + "}}" in paragraph.text for placeholder in
                                self.replacements.keys()):
                            alignment = paragraph.alignment

                            new_text = await self.replace_placeholders_in_text(paragraph.text, self.replacements)

                            paragraph.clear()
                            run = paragraph.add_run(new_text)

                            run.font.name = "Times New Roman"

                            paragraph.alignment = alignment


    async def process_document(self, document: Document, params: Dict[str, Any]) -> Document:
        """
        Обрабатывает DOCX-документ, заменяя плейсхолдеры на данные

        Args:
            document: Объект документа Word
            params: Параметры для обработки документа

        Returns:
            Обработанный документ Word
        """

        template_id = params.get('template_id')

        template = await self.get_template(template_id)
        template_type = await self.determine_template_type(template_id, template.name)

        if template_type == self.TEMPLATE_TYPE_MASTER_TITLE:
            await self.process_master_title(document, params)
        elif template_type == self.TEMPLATE_TYPE_BACHELOR_TITLE:
            await self.process_bachelor_title(document, params)
        elif template_type == self.TEMPLATE_TYPE_EXAM_TICKET:
            await self.process_exam_ticket(document, params)
        elif template_type == self.TEMPLATE_TYPE_ABSTRACT:
            await self.process_abstract(document, params)
        elif template_type == self.TEMPLATE_TYPE_LAB_WORK:
            await self.process_lab_work(document, params)
        elif template_type == self.TEMPLATE_TYPE_COURSE_WORK:
            await self.process_course_work(document, params)
        elif template_type == self.TEMPLATE_TYPE_COURSE_PROJECT:
            await self.process_course_project(document, params)
        elif template_type == self.TEMPLATE_TYPE_PRACTICE_THEMES:
            await self.process_practice_themes(document, params)
        elif template_type == self.TEMPLATE_TYPE_PUBLICATIONS:
            await self.process_publications(document, params)
        elif template_type == self.TEMPLATE_TYPE_LITERATURE:
            await self.process_literature(document, params)
        elif template_type == self.TEMPLATE_TYPE_EXAM_QUESTIONS:
            await self.process_exam_questions(document, params)
        elif template_type == self.TEMPLATE_TYPE_TEACHER_SCHEDULE:
            await self.process_teacher_schedule(document, params)
        elif template_type == self.TEMPLATE_TYPE_SECTION_PROGRAM:
            await self.process_section_program(document, params)
        elif template_type == self.TEMPLATE_TYPE_PRACTICE_REPORT:
            await self.process_practice_report(document, params)
        elif template_type == self.TEMPLATE_TYPE_TASK:
            await self.process_task(document, params)
        elif template_type == self.TEMPLATE_TYPE_CLASSROOM_SCHEDULE:
            await self.process_classroom_schedule(document, params)
        else:
            await self.generic_default_replacing(document, params)

        await self.apply_times_new_roman(document)

        return document
